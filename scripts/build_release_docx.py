from __future__ import annotations

import argparse
from datetime import datetime, timezone
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


RELEASE_TIMESTAMP = datetime(2026, 7, 15, tzinfo=timezone.utc)
RELEASE_AUTHOR = "Deep Agent"


def add_run(paragraph, text: str, *, code: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Consolas" if code else "Arial"
    run.font.size = Pt(9 if code else 10.5)
    if not code:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def shade(paragraph) -> None:
    properties = paragraph._element.get_or_add_pPr()
    value = OxmlElement("w:shd")
    value.set(qn("w:fill"), "F3F4F6")
    properties.append(value)


def markdown_to_docx(source: Path, destination: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    title = next(
        (match.group(1).strip() for line in lines if (match := re.match(r"^#\s+(.+)$", line))),
        source.stem,
    )
    properties = document.core_properties
    properties.title = title
    properties.author = RELEASE_AUTHOR
    properties.last_modified_by = RELEASE_AUTHOR
    properties.created = RELEASE_TIMESTAMP
    properties.modified = RELEASE_TIMESTAMP
    properties.revision = 1
    properties.keywords = "DeepSeek Agent V3, v0.11.0, 2026-07-15"
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size in (("Title", 22), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        styles[name].font.name = "Arial"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles[name].font.size = Pt(size)

    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        if not code_lines:
            return
        paragraph = document.add_paragraph()
        shade(paragraph)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(6)
        add_run(paragraph, "\n".join(code_lines), code=True)
        code_lines.clear()

    for line in lines:
        if line.startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            document.add_paragraph()
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            if level == 1 and len(document.paragraphs) == 0:
                document.add_heading(heading.group(2), level=0)
            else:
                document.add_heading(heading.group(2), level=level)
            continue
        ordered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if ordered:
            paragraph = document.add_paragraph()
            add_run(paragraph, f"{ordered.group(1)}. {ordered.group(2)}")
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_run(paragraph, line[2:])
            continue
        paragraph = document.add_paragraph()
        add_run(paragraph, line)

    flush_code()
    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    add_run(footer, "DeepSeek Agent V3 · generated from the versioned Markdown source")
    footer.add_run().add_break(WD_BREAK.LINE)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()
    markdown_to_docx(args.source, args.destination)


if __name__ == "__main__":
    main()
