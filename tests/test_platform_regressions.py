from __future__ import annotations

from datetime import datetime, timezone
import errno
import os
from pathlib import Path
import re
import subprocess
import sys
from types import SimpleNamespace
import tomllib
from zipfile import ZipFile

from docx import Document
import pytest

from agent import file_lock
from agent.capability_health import CapabilityHealthManager
from agent.memory import MemoryStore
from agent.project import ProjectManager
from agent.tools import ToolManager
from agent.tools.base import run_command


def test_cli_import_does_not_require_posix_fcntl() -> None:
    script = """
import builtins

original_import = builtins.__import__

def import_without_fcntl(name, *args, **kwargs):
    if name == "fcntl":
        raise ModuleNotFoundError("No module named 'fcntl'", name="fcntl")
    return original_import(name, *args, **kwargs)

builtins.__import__ = import_without_fcntl
import agent.cli
"""

    completed = subprocess.run(
        [sys.executable, "-c", script],
        check=False,
        capture_output=True,
        text=True,
        timeout=15,
    )

    assert completed.returncode == 0, completed.stderr


def test_windows_file_lock_initializes_and_locks_one_byte(tmp_path: Path, monkeypatch) -> None:
    calls: list[tuple[int, int, int]] = []
    fake_msvcrt = SimpleNamespace(
        LK_NBLCK=11,
        LK_LOCK=12,
        LK_UNLCK=13,
        locking=lambda fd, mode, count: calls.append((fd, mode, count)),
    )
    monkeypatch.setitem(sys.modules, "msvcrt", fake_msvcrt)
    monkeypatch.setattr(file_lock, "os", SimpleNamespace(name="nt", SEEK_END=os.SEEK_END))
    path = tmp_path / "windows.lock"

    with path.open("a+") as handle:
        file_lock.lock_exclusive(handle, nonblocking=True)
        file_lock.unlock(handle)
        descriptor = handle.fileno()

    assert path.read_bytes() == b"\0"
    assert calls == [(descriptor, fake_msvcrt.LK_NBLCK, 1), (descriptor, fake_msvcrt.LK_UNLCK, 1)]


def test_windows_nonblocking_lock_error_has_portable_exception(tmp_path: Path, monkeypatch) -> None:
    def unavailable(_fd: int, _mode: int, _count: int) -> None:
        raise OSError("locked")

    fake_msvcrt = SimpleNamespace(LK_NBLCK=11, LK_LOCK=12, LK_UNLCK=13, locking=unavailable)
    monkeypatch.setitem(sys.modules, "msvcrt", fake_msvcrt)
    monkeypatch.setattr(file_lock, "os", SimpleNamespace(name="nt", SEEK_END=os.SEEK_END))

    with (tmp_path / "windows.lock").open("a+") as handle:
        with pytest.raises(file_lock.FileLockUnavailable, match="already held"):
            file_lock.lock_exclusive(handle, nonblocking=True)


def test_posix_nonblocking_lock_maps_eacces_to_portable_exception(tmp_path: Path, monkeypatch) -> None:
    def unavailable(_fd: int, _flags: int) -> None:
        raise PermissionError(errno.EACCES, "locked")

    fake_fcntl = SimpleNamespace(LOCK_EX=1, LOCK_NB=2, flock=unavailable)
    monkeypatch.setitem(sys.modules, "fcntl", fake_fcntl)
    monkeypatch.setattr(file_lock, "os", SimpleNamespace(name="posix"))

    with (tmp_path / "posix.lock").open("a+") as handle:
        with pytest.raises(file_lock.FileLockUnavailable, match="already held"):
            file_lock.lock_exclusive(handle, nonblocking=True)


@pytest.mark.skipif(os.name == "nt", reason="shell=True uses the platform shell; this checks POSIX quoting")
def test_shell_mode_preserves_metacharacters_inside_one_argument(tmp_path: Path) -> None:
    marker = tmp_path / "must-not-exist"
    payload = f"literal; touch {marker}"

    result = run_command(["printf", "%s", payload], cwd=tmp_path, timeout=5, shell=True)

    assert result.success is True
    assert result.stdout == payload
    assert not marker.exists()


def test_tool_manager_uses_injected_capability_health(tmp_path: Path, make_config) -> None:
    root = tmp_path / "project"
    root.mkdir()
    config = make_config()
    project = ProjectManager(config).resolve_project(root)
    memory = MemoryStore(config)
    memory.sync_project(project)
    injected = CapabilityHealthManager(config, "injected-health")

    manager = ToolManager(config, project, memory, health=injected)

    assert manager.health is injected


def test_core_requirements_match_project_dependencies() -> None:
    root = Path(__file__).resolve().parents[1]
    project = tomllib.loads((root / "pyproject.toml").read_text(encoding="utf-8"))["project"]
    requirements_text = (root / "requirements.txt").read_text(encoding="utf-8")
    requirements = [
        line for raw_line in requirements_text.splitlines() if (line := raw_line.strip()) and not line.startswith("#")
    ]
    requirement_names = {line.split(">=", maxsplit=1)[0].lower() for line in requirements}

    assert requirements == project["dependencies"]
    assert requirement_names.isdisjoint({"chromadb", "playwright"})
    assert "pip install -e '.[browser,vector,semantic,document]'" in requirements_text


def test_ruff_version_and_lint_contract_are_release_pinned() -> None:
    root = Path(__file__).resolve().parents[1]
    configuration = tomllib.loads((root / "pyproject.toml").read_text(encoding="utf-8"))

    assert "ruff==0.15.21" in configuration["project"]["optional-dependencies"]["dev"]
    assert configuration["tool"]["ruff"]["required-version"] == "==0.15.21"
    assert configuration["tool"]["ruff"]["lint"]["select"] == ["E4", "E7", "E9", "F"]


def test_release_tree_exposes_only_current_user_documents() -> None:
    root = Path(__file__).resolve().parents[1]
    configuration = tomllib.loads((root / "pyproject.toml").read_text(encoding="utf-8"))
    version = configuration["project"]["version"]
    expected_root_words = {
        f"DeepSeek-Agent-V3-使用说明-{version}.docx",
        f"DeepSeek-Agent-V3-工作日志-{version}.docx",
    }
    actual_root_words = {path.name for path in root.glob("*.docx")}

    user_docs = root / "user-docs"
    expected_user_docs = {
        Path("DeepSeek-Agent-V3-使用说明.md"),
        Path("DeepSeek-Agent-V3-工作日志.md"),
        *(Path(name) for name in expected_root_words),
    }
    actual_user_docs = {path.relative_to(user_docs) for path in user_docs.rglob("*") if path.is_file()}
    unexpected_directories = {path.relative_to(user_docs) for path in user_docs.rglob("*") if path.is_dir()}

    assert actual_root_words == expected_root_words
    assert actual_user_docs == expected_user_docs
    assert not unexpected_directories
    assert not any(path.name == "AGENTS.md" for path in actual_user_docs)


def test_current_release_word_files_reopen_with_matching_sources_and_metadata(tmp_path: Path) -> None:
    root = Path(__file__).resolve().parents[1]
    configuration = tomllib.loads((root / "pyproject.toml").read_text(encoding="utf-8"))
    version = configuration["project"]["version"]
    user_docs = root / "user-docs"
    documents = (
        ("使用说明", "更新日期"),
        ("工作日志", "日期"),
    )

    for label, date_label in documents:
        markdown_path = user_docs / f"DeepSeek-Agent-V3-{label}.md"
        word_name = f"DeepSeek-Agent-V3-{label}-{version}.docx"
        root_word = root / word_name
        user_word = user_docs / word_name
        markdown = markdown_path.read_text(encoding="utf-8")
        title_match = re.search(r"^#\s+(.+)$", markdown, flags=re.MULTILINE)
        date_match = re.search(rf"^{date_label}：([0-9]{{4}}-[0-9]{{2}}-[0-9]{{2}})$", markdown, flags=re.MULTILINE)

        assert title_match is not None
        assert date_match is not None
        release_date = date_match.group(1)
        expected_timestamp = datetime.strptime(release_date, "%Y-%m-%d").replace(tzinfo=timezone.utc)
        assert root_word.read_bytes() == user_word.read_bytes()

        regenerated = tmp_path / word_name
        completed = subprocess.run(
            [
                sys.executable,
                str(root / "scripts" / "build_release_docx.py"),
                str(markdown_path),
                str(regenerated),
                "--version",
                version,
                "--release-date",
                release_date,
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=30,
        )
        assert completed.returncode == 0, completed.stderr

        with ZipFile(root_word) as archive:
            assert archive.testzip() is None
            assert {
                "[Content_Types].xml",
                "_rels/.rels",
                "docProps/core.xml",
                "word/document.xml",
            }.issubset(archive.namelist())

        reopened = Document(root_word)
        regenerated_document = Document(regenerated)
        properties = reopened.core_properties
        assert properties.title == title_match.group(1).strip()
        assert properties.author == "Deep Agent"
        assert properties.last_modified_by == "Deep Agent"
        assert properties.created == expected_timestamp
        assert properties.modified == expected_timestamp
        assert properties.version == version
        assert properties.keywords == f"DeepSeek Agent V3, v{version}, {release_date}"
        assert any(paragraph.text.strip() for paragraph in reopened.paragraphs)
        assert [paragraph.text for paragraph in reopened.paragraphs] == [
            paragraph.text for paragraph in regenerated_document.paragraphs
        ]
        assert [paragraph.text for paragraph in reopened.sections[0].footer.paragraphs] == [
            paragraph.text for paragraph in regenerated_document.sections[0].footer.paragraphs
        ]
        assert version in markdown
