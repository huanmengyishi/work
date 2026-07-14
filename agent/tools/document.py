from __future__ import annotations

import shutil
import json
import sys
import tempfile
from io import BytesIO
from pathlib import Path

from .base import (
    DEFAULT_MAX_RESULT_SOURCE_BYTES,
    BoundedByteCapture,
    ToolResult,
    bound_text_source,
    bounded_result_source_bytes,
    run_command,
)
from .pathsafe import resolve_project_path


TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".go",
    ".rs",
    ".gd",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".xml",
    ".html",
    ".css",
    ".scss",
    ".sh",
    ".sql",
}

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
WORD_SUFFIXES = {".docx", ".doc", ".odt", ".rtf"}
AI_TOOLS_PARSER = Path.home() / ".local" / "share" / "ai-tools" / "app" / "parser.py"
AI_TOOLS_LAUNCHER = Path.home() / ".local" / "bin" / "ai-parser"


class DocumentTool:
    def __init__(
        self,
        cwd: Path,
        timeout: int = 180,
        max_input_bytes: int = 25_000_000,
        max_result_bytes: int = DEFAULT_MAX_RESULT_SOURCE_BYTES,
    ) -> None:
        self.cwd = cwd
        self.timeout = timeout
        self.max_input_bytes = max(1, min(int(max_input_bytes), 100_000_000))
        self.max_result_bytes = bounded_result_source_bytes(max_result_bytes)

    def parse(self, path: str, ocr: bool = True) -> ToolResult:
        file_path = self._resolve(path)
        if not file_path.exists():
            return ToolResult(False, "", f"file not found: {file_path}")
        if file_path.stat().st_size > self.max_input_bytes:
            return ToolResult(False, "", f"document exceeds {self.max_input_bytes} bytes")
        suffix = file_path.suffix.lower()
        if suffix in {".pdf", *IMAGE_SUFFIXES, *WORD_SUFFIXES}:
            ai_tools_result = self._parse_with_ai_tools(file_path)
            if ai_tools_result.ok:
                return ai_tools_result
        if suffix in TEXT_SUFFIXES:
            return self._parse_text(file_path)
        if suffix == ".pdf":
            return self._parse_pdf(file_path, ocr=ocr)
        if suffix in IMAGE_SUFFIXES:
            return self._parse_image(file_path)
        if suffix in WORD_SUFFIXES:
            return ToolResult(False, "", "word parsing requires ai-parser dependencies or pandoc")
        return ToolResult(False, "", f"unsupported document type: {suffix or '<none>'}")

    def render_docx(self, *, title: str, markdown: str) -> tuple[bytes, dict[str, object]]:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:
            raise RuntimeError("Word generation requires the python-docx package") from exc
        document = Document()
        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)
        if title.strip():
            document.add_heading(title.strip()[:500], level=0)
        paragraph_count = 0
        heading_count = 0
        for raw_line in str(markdown).splitlines():
            line = raw_line.rstrip()
            if not line:
                document.add_paragraph()
                paragraph_count += 1
                continue
            stripped = line.lstrip()
            hashes = len(stripped) - len(stripped.lstrip("#"))
            if hashes and hashes <= 6 and stripped[hashes:].startswith(" "):
                document.add_heading(stripped[hashes + 1 :].strip()[:1000], level=min(hashes, 9))
                heading_count += 1
                continue
            if stripped.startswith(("- ", "* ")):
                document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            elif len(stripped) > 3 and stripped[0].isdigit() and ". " in stripped[:5]:
                document.add_paragraph(stripped.split(". ", 1)[1], style="List Number")
            else:
                document.add_paragraph(line)
            paragraph_count += 1
        stream = BytesIO()
        document.save(stream)
        content = stream.getvalue()
        return content, {
            "format": "docx",
            "bytes": len(content),
            "paragraph_count": paragraph_count,
            "heading_count": heading_count,
        }

    def _resolve(self, path: str) -> Path:
        return resolve_project_path(self.cwd, path, require_file=True)

    def _parse_text(self, file_path: Path) -> ToolResult:
        try:
            text = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = file_path.read_text(encoding="utf-8", errors="replace")
        return self._markdown_result(file_path, text, {"path": str(file_path), "parser": "text"})

    def _parse_pdf(self, file_path: Path, ocr: bool) -> ToolResult:
        if shutil.which("pdftotext"):
            result = run_command(
                ["pdftotext", "-layout", str(file_path), "-"],
                cwd=self.cwd,
                timeout=self.timeout,
                max_output_bytes=self.max_result_bytes,
            )
            if result.ok and result.output.strip():
                return self._markdown_result(
                    file_path,
                    result.output,
                    {"path": str(file_path), "parser": "pdftotext", **_source_metadata(result.data)},
                )
        if ocr:
            return self._ocr_pdf(file_path)
        return ToolResult(False, "", "PDF parsing failed and OCR is disabled")

    def _parse_with_ai_tools(self, file_path: Path) -> ToolResult:
        if not AI_TOOLS_LAUNCHER.exists() and not AI_TOOLS_PARSER.exists():
            return ToolResult(False, "", f"ai-parser not found: {AI_TOOLS_PARSER}")
        with tempfile.TemporaryDirectory(prefix="deep-agent-ai-parser-") as tmp:
            output_dir = Path(tmp)
            command = [str(AI_TOOLS_LAUNCHER)] if AI_TOOLS_LAUNCHER.exists() else [sys.executable, str(AI_TOOLS_PARSER)]
            result = run_command(
                [*command, str(file_path), "-o", str(output_dir)],
                cwd=self.cwd,
                timeout=self.timeout,
                max_output_bytes=self.max_result_bytes,
            )
            if not result.ok:
                return result
            json_files = sorted(output_dir.glob("*.json"))
            text_files = sorted(output_dir.glob("*.txt"))
            metadata = {"path": str(file_path), "parser": "ai-parser"}
            text = ""
            if json_files:
                try:
                    if json_files[0].stat().st_size > self.max_input_bytes:
                        return ToolResult(
                            False,
                            "",
                            f"ai-parser JSON output exceeds {self.max_input_bytes} bytes",
                            data={
                                **metadata,
                                "source_truncated": True,
                                "source_original_bytes": json_files[0].stat().st_size,
                                "source_original_bytes_known": True,
                                "source_captured_bytes": 0,
                            },
                        )
                    payload = json.loads(json_files[0].read_text(encoding="utf-8"))
                    text = str(payload.get("text") or "")
                    if isinstance(payload.get("metadata"), dict):
                        metadata.update(payload["metadata"])
                    metadata["kind"] = payload.get("kind")
                except json.JSONDecodeError:
                    text = ""
            if not text and text_files:
                if text_files[0].stat().st_size > self.max_input_bytes:
                    return ToolResult(
                        False,
                        "",
                        f"ai-parser text output exceeds {self.max_input_bytes} bytes",
                        data={
                            **metadata,
                            "source_truncated": True,
                            "source_original_bytes": text_files[0].stat().st_size,
                            "source_original_bytes_known": True,
                            "source_captured_bytes": 0,
                        },
                    )
                text = text_files[0].read_text(encoding="utf-8", errors="replace")
            if not text.strip():
                return ToolResult(False, result.output, "ai-parser produced no text", data=metadata)
            return self._markdown_result(file_path, text, metadata)

    def _parse_image(self, file_path: Path) -> ToolResult:
        if not shutil.which("tesseract"):
            return ToolResult(False, "", "tesseract is not installed")
        result = run_command(
            ["tesseract", str(file_path), "stdout", "-l", "eng+chi_sim"],
            cwd=self.cwd,
            timeout=self.timeout,
            max_output_bytes=self.max_result_bytes,
        )
        if not result.ok:
            result = run_command(
                ["tesseract", str(file_path), "stdout"],
                cwd=self.cwd,
                timeout=self.timeout,
                max_output_bytes=self.max_result_bytes,
            )
        if not result.ok:
            return result
        return self._markdown_result(
            file_path,
            result.output,
            {"path": str(file_path), "parser": "tesseract", **_source_metadata(result.data)},
        )

    def _ocr_pdf(self, file_path: Path) -> ToolResult:
        if not shutil.which("magick") and not shutil.which("convert"):
            return ToolResult(False, "", "PDF OCR requires ImageMagick magick/convert")
        if not shutil.which("tesseract"):
            return ToolResult(False, "", "PDF OCR requires tesseract")
        converter = shutil.which("magick") or shutil.which("convert")
        assert converter is not None
        with tempfile.TemporaryDirectory(prefix="deep-agent-ocr-") as tmp:
            tmpdir = Path(tmp)
            out_pattern = tmpdir / "page.png"
            cmd = [converter]
            if Path(converter).name == "magick":
                cmd.extend([str(file_path), str(out_pattern)])
            else:
                cmd.extend(["-density", "180", str(file_path), str(out_pattern)])
            convert_result = run_command(
                cmd,
                cwd=self.cwd,
                timeout=self.timeout,
                max_output_bytes=self.max_result_bytes,
            )
            if not convert_result.ok:
                return convert_result
            capture = BoundedByteCapture(self.max_result_bytes)
            found_text = False
            for image in sorted(tmpdir.glob("page*.png")):
                ocr_result = self._parse_image(image)
                if ocr_result.ok and ocr_result.output.strip():
                    found_text = True
                    capture.feed(f"## {image.name}\n\n{ocr_result.output}\n\n".encode("utf-8"))
            if not found_text:
                return ToolResult(False, "", "OCR produced no text")
            return self._markdown_result(
                file_path,
                capture.text(),
                {"path": str(file_path), "parser": "pdf-ocr", **capture.metadata()},
            )

    def _markdown_result(self, path: Path, content: str, data: dict[str, object]) -> ToolResult:
        rendered, bounded_metadata = bound_text_source(
            markdown_document(path, content),
            self.max_result_bytes,
        )
        upstream = _source_metadata(data)
        if upstream.get("source_truncated"):
            bounded_metadata.update(upstream)
        merged = dict(data)
        merged.update(bounded_metadata)
        return ToolResult(True, rendered, data=merged)


def markdown_document(path: Path, content: str) -> str:
    return f"# Parsed Document\n\n- Source: `{path}`\n\n---\n\n{content.strip()}\n"


def _source_metadata(value: object) -> dict[str, object]:
    if not isinstance(value, dict):
        return {}
    return {str(key): item for key, item in value.items() if str(key).startswith("source_")}
